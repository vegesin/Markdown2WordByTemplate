#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "o": "urn:schemas-microsoft-com:office:office",
}
INLINE_RE = re.compile(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)", re.DOTALL)
DISPLAY_RE = re.compile(r"\\\[(.+?)\\\]", re.DOTALL)


def paragraph_texts(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text


def main() -> int:
    parser = argparse.ArgumentParser(description="Audit a stage-1 DOCX before MathType conversion.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--expect-inline", type=int)
    parser.add_argument("--expect-display", type=int)
    parser.add_argument("--expect-visio", type=int)
    args = parser.parse_args()

    path = args.docx.resolve()
    doc = Document(path)
    texts = list(paragraph_texts(doc))
    display = sum(len(DISPLAY_RE.findall(text)) for text in texts)
    inline = sum(len(INLINE_RE.findall(DISPLAY_RE.sub("", text))) for text in texts)
    visio = sum(len(re.findall(r"\[VISIO:diagram-\d{2}\]", text)) for text in texts)
    headings: dict[str, int] = {}
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading "):
            headings[paragraph.style.name] = headings.get(paragraph.style.name, 0) + 1

    with zipfile.ZipFile(path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
        table_borders = []
        for table in root.xpath(".//w:tbl", namespaces=NS):
            borders = table.find("w:tblPr/w:tblBorders", NS)
            values = {}
            if borders is not None:
                for edge in borders:
                    values[etree.QName(edge).localname] = edge.get(f"{{{NS['w']}}}val")
            header_bottom = table.find(".//w:tr[1]/w:tc[1]/w:tcPr/w:tcBorders/w:bottom", NS)
            table_borders.append({
                "borders": values,
                "header_bottom": header_bottom.get(f"{{{NS['w']}}}val") if header_bottom is not None else None,
            })
        formula_paragraphs = {"inline_total": 0, "inline_at_least": 0, "display_total": 0, "display_at_least": 0}
        for paragraph in root.xpath(".//w:p", namespaces=NS):
            text = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
            spacing = paragraph.find("w:pPr/w:spacing", NS)
            rule = spacing.get(f"{{{NS['w']}}}lineRule") if spacing is not None else None
            line = int(spacing.get(f"{{{NS['w']}}}line", "0")) if spacing is not None else 0
            if DISPLAY_RE.search(text):
                formula_paragraphs["display_total"] += 1
                if rule == "atLeast" and line >= 600:
                    formula_paragraphs["display_at_least"] += 1
            elif INLINE_RE.search(text):
                formula_paragraphs["inline_total"] += 1
                if rule == "atLeast" and line >= 500:
                    formula_paragraphs["inline_at_least"] += 1
        summary = {
            "path": str(path),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_latex": inline,
            "display_latex": display,
            "visio_placeholders": visio,
            "headings": headings,
            "omml_nodes": len(root.xpath(".//m:oMath | .//m:oMathPara", namespaces=NS)),
            "ole_objects": len(root.xpath(".//w:object | .//o:OLEObject", namespaces=NS)),
            "equation_seq_fields": len(root.xpath(".//w:fldSimple[contains(@w:instr, 'SEQ Equation')]", namespaces=NS)),
            "header_parts": len([name for name in zf.namelist() if name.startswith("word/header")]),
            "footer_parts": len([name for name in zf.namelist() if name.startswith("word/footer")]),
            "table_borders": table_borders,
            "formula_paragraph_spacing": formula_paragraphs,
        }
    print(json.dumps(summary, ensure_ascii=False, indent=2))

    expected = ((args.expect_inline, inline), (args.expect_display, display), (args.expect_visio, visio))
    if any(want is not None and want != actual for want, actual in expected):
        return 2
    if summary["omml_nodes"] or summary["ole_objects"]:
        return 3
    if summary["equation_seq_fields"] != 0:
        return 4
    spacing = summary["formula_paragraph_spacing"]
    if spacing["display_total"] != spacing["display_at_least"]:
        return 6
    if spacing["inline_total"] != spacing["inline_at_least"]:
        return 6
    for table in table_borders:
        borders = table["borders"]
        if borders.get("top") != "single" or borders.get("bottom") != "single":
            return 5
        if any(borders.get(edge) != "nil" for edge in ("left", "right", "insideH", "insideV")):
            return 5
        if table["header_bottom"] != "single":
            return 5
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
