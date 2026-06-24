#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


MERMAID_RE = re.compile(r"^```mermaid\s*\r?\n(.*?)^```\s*$", re.MULTILINE | re.DOTALL)
INLINE_MATH_RE = re.compile(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)", re.DOTALL)
INLINE_PAREN_RE = re.compile(r"\\\((.+?)\\\)", re.DOTALL)
DISPLAY_DOLLAR_RE = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
DISPLAY_BRACKET_RE = re.compile(r"\\\[(.+?)\\\]", re.DOTALL)
FIGURE_LABEL = "\u56fe"
TABLE_LABEL = "\u8868"
CHINESE_NUMERAL_RE = "\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u5343\u4e07"
CAPTION_SEPARATORS_RE = r"[\.\u3001\uff0e:：\s-]*"


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Build a template-based DOCX with LaTeX preserved for MathType.")
    p.add_argument("markdown", type=Path)
    p.add_argument("template_or_output", type=Path)
    p.add_argument("output", type=Path, nargs="?")
    p.add_argument("--pandoc", default="pandoc")
    p.add_argument("--visio-dir", type=Path)
    p.add_argument(
        "--default-template",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "assets" / "default-template.docx",
    )
    return p.parse_args()


def source_math_counts(text: str) -> tuple[int, int]:
    without_code = MERMAID_RE.sub("", text)
    displays = DISPLAY_DOLLAR_RE.findall(without_code) + DISPLAY_BRACKET_RE.findall(without_code)
    stripped = DISPLAY_DOLLAR_RE.sub("", without_code)
    stripped = DISPLAY_BRACKET_RE.sub("", stripped)
    return len(INLINE_MATH_RE.findall(stripped)) + len(INLINE_PAREN_RE.findall(stripped)), len(displays)


def extract_mermaid(text: str, output_dir: Path) -> tuple[str, list[dict[str, str]]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest: list[dict[str, str]] = []

    def replace(match: re.Match[str]) -> str:
        index = len(manifest) + 1
        diagram_id = f"diagram-{index:02d}"
        source_name = f"{diagram_id}.mmd"
        source = match.group(1).strip() + "\n"
        (output_dir / source_name).write_text(source, encoding="utf-8")
        manifest.append({
            "id": diagram_id,
            "placeholder": f"[VISIO:{diagram_id}]",
            "source": source_name,
            "diagram_type": source.splitlines()[0] if source.splitlines() else "",
        })
        return f"\n\n[VISIO:{diagram_id}]\n\n"

    replaced = MERMAID_RE.sub(replace, text)
    (output_dir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return replaced, manifest


def normalize_display_math_blocks(text: str) -> str:
    def dollar(match: re.Match[str]) -> str:
        formula = re.sub(r"\s+", " ", match.group(1)).strip()
        return f"\n\n\\[{formula}\\]\n\n"

    def bracket(match: re.Match[str]) -> str:
        formula = re.sub(r"\s+", " ", match.group(1)).strip()
        return f"\n\n\\[{formula}\\]\n\n"

    return DISPLAY_BRACKET_RE.sub(bracket, DISPLAY_DOLLAR_RE.sub(dollar, text))


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def first_existing_style(doc: Document, names: list[str]):
    by_name = {style.name: style for style in doc.styles}
    for name in names:
        if name in by_name:
            return by_name[name]
    return None


def ensure_caption_style(doc: Document):
    try:
        return doc.styles["Caption"]
    except KeyError:
        style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "宋体"
        style.font.size = Pt(10.5)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)
        return style


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='drawing']"))


def drawing_alt_text(paragraph) -> str:
    texts: list[str] = []
    for node in paragraph._p.xpath(".//*[local-name()='docPr']"):
        for key in ("descr", "title"):
            value = node.get(key)
            if value:
                value = value.strip()
                if value and value not in texts:
                    texts.append(value)
    return " ".join(texts).strip()


def is_visio_placeholder(paragraph) -> bool:
    return bool(re.fullmatch(r"\s*\[VISIO:diagram-\d{2}\]\s*", paragraph.text.strip()))


def caption_text_body(text: str, label: str) -> str | None:
    pattern = rf"^\s*{re.escape(label)}\s*(?:\d+|[{CHINESE_NUMERAL_RE}]+)?{CAPTION_SEPARATORS_RE}(.*?)\s*$"
    match = re.match(pattern, text)
    if not match:
        return None
    body = clean_caption_body(match.group(1), label)
    return body or None


def clean_caption_body(text: str, label: str) -> str:
    body = re.sub(r"\s+", " ", text or "").strip()
    if not body:
        return ""
    # Remove an existing Markdown/alt-text caption prefix such as:
    # "图 1 系统框架", "表1: 指标", or the Pandoc image-alt variant
    # "1 系统框架". The Word SEQ field supplies the live number.
    with_label = rf"^\s*{re.escape(label)}\s*(?:\d+|[{CHINESE_NUMERAL_RE}]+)?{CAPTION_SEPARATORS_RE}"
    body = re.sub(with_label, "", body).strip()
    bare_number = rf"^\s*(?:\d+|[{CHINESE_NUMERAL_RE}]+){CAPTION_SEPARATORS_RE}"
    body = re.sub(bare_number, "", body).strip()
    return body


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return paragraph._parent.paragraphs[[p._p for p in paragraph._parent.paragraphs].index(new_p)]


def insert_paragraph_before_table(table):
    new_p = OxmlElement("w:p")
    table._tbl.addprevious(new_p)
    body = table._parent
    return body.paragraphs[[p._p for p in body.paragraphs].index(new_p)]


def write_caption(paragraph, label: str, body: str, number: int) -> None:
    body = clean_caption_body(body, label)
    clear_paragraph(paragraph)
    paragraph.style = paragraph.part.document.styles["Caption"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = label == TABLE_LABEL
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run(label)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f"SEQ {label} \\* ARABIC")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = str(number)
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    default_body = "\u8bf7\u586b\u5199\u56fe\u9898" if label == FIGURE_LABEL else "\u8bf7\u586b\u5199\u8868\u9898"
    paragraph.add_run(" " + (body.strip() if body.strip() else default_body))


def iter_body_blocks(doc: Document):
    paragraphs_by_element = {p._p: p for p in doc.paragraphs}
    tables_by_element = {t._tbl: t for t in doc.tables}
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p") and child in paragraphs_by_element:
            yield ("paragraph", paragraphs_by_element[child])
        elif child.tag == qn("w:tbl") and child in tables_by_element:
            yield ("table", tables_by_element[child])


def apply_captions(doc: Document) -> dict[str, int]:
    ensure_caption_style(doc)
    blocks = list(iter_body_blocks(doc))
    figure_no = 0
    table_no = 0
    figure_captions = 0
    table_captions = 0
    inserted_figures = 0
    inserted_tables = 0

    for index, (kind, block) in enumerate(blocks):
        if kind != "paragraph" or not (paragraph_has_drawing(block) or is_visio_placeholder(block)):
            continue
        paragraph = block
        figure_no += 1
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        body = drawing_alt_text(paragraph)
        caption_paragraph = None
        if index + 1 < len(blocks) and blocks[index + 1][0] == "paragraph":
            candidate = blocks[index + 1][1]
            parsed = caption_text_body(candidate.text, FIGURE_LABEL)
            if parsed is not None:
                caption_paragraph = candidate
                body = parsed
        if caption_paragraph is None:
            caption_paragraph = insert_paragraph_after(paragraph)
            inserted_figures += 1
        write_caption(caption_paragraph, FIGURE_LABEL, body, figure_no)
        figure_captions += 1

    blocks = list(iter_body_blocks(doc))
    for index, (kind, block) in enumerate(blocks):
        if kind != "table":
            continue
        table = block
        table_no += 1
        body = ""
        caption_paragraph = None
        if index > 0 and blocks[index - 1][0] == "paragraph":
            candidate = blocks[index - 1][1]
            parsed = caption_text_body(candidate.text, TABLE_LABEL)
            if parsed is not None:
                caption_paragraph = candidate
                body = parsed
        if caption_paragraph is None:
            caption_paragraph = insert_paragraph_before_table(table)
            inserted_tables += 1
        write_caption(caption_paragraph, TABLE_LABEL, body, table_no)
        table_captions += 1

    return {
        "figure_captions": figure_captions,
        "table_captions": table_captions,
        "inserted_figure_caption_placeholders": inserted_figures,
        "inserted_table_caption_placeholders": inserted_tables,
    }


def set_line_spacing_at_least(paragraph, twips: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:line"), str(twips))
    spacing.set(qn("w:lineRule"), "atLeast")


def format_display_equations(doc: Document) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        match = re.fullmatch(r"\s*\\\[(.*?)\\\]\s*", paragraph.text, re.DOTALL)
        if not match:
            continue
        count += 1
        formula = re.sub(r"\s+", " ", match.group(1)).strip()
        clear_paragraph(paragraph)
        display_style = first_existing_style(doc, ["MTDisplayEquation", "MTDisplayE"])
        if display_style is not None:
            paragraph.style = display_style
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        set_line_spacing_at_least(paragraph, 600)
        paragraph.add_run(f"\\[{formula}\\]")
    return count


def all_content_paragraphs(doc: Document):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def format_inline_formula_spacing(doc: Document) -> int:
    count = 0
    for paragraph in all_content_paragraphs(doc):
        text = DISPLAY_BRACKET_RE.sub("", paragraph.text)
        if INLINE_MATH_RE.search(text):
            count += 1
            set_line_spacing_at_least(paragraph, 500)
    return count


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if edge in ("top", "bottom"):
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "12")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)

    if table.rows:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            old_tc = tc_pr.find(qn("w:tcBorders"))
            if old_tc is not None:
                tc_pr.remove(old_tc)
            tc_borders = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), "000000")
            tc_borders.append(bottom)
            tc_pr.append(tc_borders)


def format_tables(doc: Document) -> int:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table)
        if table.rows:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or col_index > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.size = Pt(12)
                        if row_index == 0:
                            run.bold = True
    return len(doc.tables)


def format_visio_placeholders(doc: Document) -> int:
    count = 0
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if re.fullmatch(r"\[VISIO:diagram-\d{2}\]", paragraph.text.strip()):
            count += 1
            if "pic" in [s.name for s in doc.styles]:
                paragraph.style = doc.styles["pic"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                run.bold = True
            if index + 1 < len(paragraphs) and re.match(r"^图\s*\d", paragraphs[index + 1].text.strip()):
                if "Caption" in [s.name for s in doc.styles]:
                    paragraphs[index + 1].style = doc.styles["Caption"]
    return count


def count_headings(doc: Document) -> dict[str, int]:
    result: dict[str, int] = {}
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading "):
            result[paragraph.style.name] = result.get(paragraph.style.name, 0) + 1
    return result


def audit_docx(path: Path) -> dict[str, int]:
    with zipfile.ZipFile(path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
        ns = {
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "o": "urn:schemas-microsoft-com:office:office",
        }
        return {
            "omml_nodes": len(root.xpath(".//m:oMath | .//m:oMathPara", namespaces=ns)),
            "ole_objects": len(root.xpath(".//w:object | .//o:OLEObject", namespaces=ns)),
            "equation_seq_fields": len(root.xpath(".//w:fldSimple[contains(@w:instr, 'SEQ Equation')]", namespaces=ns)),
        }


def main() -> int:
    args = parse_args()
    markdown = args.markdown.resolve()
    if args.output is None:
        template = args.default_template.resolve()
        output = args.template_or_output.resolve()
    else:
        template = args.template_or_output.resolve()
        output = args.output.resolve()
    if not markdown.is_file():
        raise FileNotFoundError(f"Markdown path does not exist: {markdown}")
    if not template.is_file():
        raise FileNotFoundError(f"Word template path does not exist: {template}")
    output.parent.mkdir(parents=True, exist_ok=True)
    visio_dir = (args.visio_dir or output.with_name(output.stem + "_visio_sources")).resolve()
    if visio_dir.exists():
        shutil.rmtree(visio_dir)

    source = markdown.read_text(encoding="utf-8")
    source_inline, source_display = source_math_counts(source)
    replaced, manifest = extract_mermaid(source, visio_dir)
    replaced = normalize_display_math_blocks(replaced)
    lua_filter = Path(__file__).with_name("preserve_math.lua")

    with tempfile.TemporaryDirectory(prefix="md-template-docx-") as tmp:
        tmp_path = Path(tmp)
        temp_md = tmp_path / markdown.name
        temp_docx = tmp_path / "pandoc.docx"
        temp_md.write_text(replaced, encoding="utf-8")
        cmd = [
            args.pandoc,
            "--from=markdown+tex_math_dollars+tex_math_single_backslash",
            "--to=docx",
            f"--reference-doc={template}",
            f"--lua-filter={lua_filter}",
            f"--resource-path={markdown.parent}",
            "--wrap=none",
            f"--output={temp_docx}",
            str(temp_md),
        ]
        subprocess.run(cmd, check=True)
        doc = Document(temp_docx)
        display_count = format_display_equations(doc)
        table_count = format_tables(doc)
        visio_count = format_visio_placeholders(doc)
        caption_summary = apply_captions(doc)
        inline_formula_paragraphs = format_inline_formula_spacing(doc)
        doc.core_properties.title = markdown.stem
        doc.core_properties.subject = "Stage-1 DOCX with LaTeX placeholders for MathType"
        doc.save(output)

    reopened = Document(output)
    inline_count = 0
    display_text_count = 0
    for p in reopened.paragraphs:
        display_text_count += len(DISPLAY_BRACKET_RE.findall(p.text))
        inline_count += len(INLINE_MATH_RE.findall(DISPLAY_BRACKET_RE.sub("", p.text)))
    for table in reopened.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    display_text_count += len(DISPLAY_BRACKET_RE.findall(p.text))
                    inline_count += len(INLINE_MATH_RE.findall(DISPLAY_BRACKET_RE.sub("", p.text)))

    package_audit = audit_docx(output)
    summary = {
        "output": str(output),
        "visio_dir": str(visio_dir),
        "source_inline_math": source_inline,
        "source_display_math": source_display,
        "docx_inline_latex": inline_count,
        "docx_display_latex": display_text_count,
        "display_paragraphs_formatted": display_count,
        "inline_formula_paragraphs_at_least_25pt": inline_formula_paragraphs,
        "tables_formatted": table_count,
        **caption_summary,
        "mermaid_sources": len(manifest),
        "visio_placeholders": visio_count,
        "headings": count_headings(reopened),
        **package_audit,
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))

    if source_inline != inline_count or source_display != display_text_count:
        raise RuntimeError("Formula counts changed during conversion")
    if len(manifest) != visio_count:
        raise RuntimeError("Mermaid placeholder count mismatch")
    if package_audit["omml_nodes"] != 0 or package_audit["ole_objects"] != 0:
        raise RuntimeError("Stage-1 DOCX unexpectedly contains OMML or OLE objects")
    if package_audit["equation_seq_fields"] != 0:
        raise RuntimeError("Stage-1 DOCX must not contain equation-number fields before MathType conversion")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
